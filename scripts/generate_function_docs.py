from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_DIR = Path("/Users/leo/Library/Mobile Documents/com~apple~CloudDocs/www/smart-reader/outputs/function-docs")


SIMPLE_ROWS = [
    ("阅读端", "账号与个人中心", "登录、注册、验证码登录、忘记密码、重置密码、学校/账号/姓名信息维护、昵称头像维护、手机号密码修改、用户反馈"),
    ("阅读端", "首页与教材浏览", "首页推荐教材、教材封面/书名/作者/出版社/ISBN展示、数字出版信息、活动展示、教材搜索、分类浏览、教材详情、教材预览"),
    ("阅读端", "教材获取与书架", "在线购买、兑换码激活、免费试读、价格展示、我的书架、教材快捷入口、班级切换"),
    ("阅读端", "智慧阅读", "正文阅读、目录导航、文内搜索、字号切换、字体切换、阅读背景切换、笔记、高亮、气泡注释、超链接、AI划线解读"),
    ("阅读端", "富媒体与资源阅读", "图片、图组、视频、音频、附件、3D模型嵌入阅读，多格式资源下载与在线预览"),
    ("阅读端", "互动答题与AI评测", "选择题、填空题、判断题、简答题，自动批阅、答案解析、错题思路、AI评价、知识点分析、学习建议、重复做题、学习追踪"),
    ("阅读端", "视听模式", "视频学习、音频学习、学习要点展示、播放控制、上下切换、全屏播放"),
    ("阅读端", "任务模式与任务工单", "任务式阅读、阅读任务、答题任务、实践任务、任务进度跟踪、子任务拆分、任务工单创建、导入、发布、撤回、跟踪"),
    ("阅读端", "知识图谱与学情图谱", "知识点图谱、知识关联、节点查看、知识点统计、知识掌握展示、阅读进度、习题进度、个性化学习路径"),
    ("阅读端", "实训与交互案例", "实训任务、交互案例、实践步骤引导、在线评分评测"),
    ("阅读端", "AI伴学", "教材专属AI问答、章节推荐、拓展知识推荐、引导式问答"),
    ("阅读端", "组群与班级教务", "创建群组、加入群组、邀请码/二维码分享、班级教材组织、群组统计、多班级管理、学生管理、班级信息维护"),
    ("阅读端", "教学模式", "教材、教案、课件、教学资料预览，课件上课、翻页教学、全屏上课、教材二分屏展示、分屏调节"),
    ("阅读端", "教学辅助工具", "分组PK、画笔、多色标注、倒计时、随机抽问、喝彩、课堂积分、小组排行、学生名单导入、AI助教"),
    ("阅读端", "智能云讲义", "课时封面、讲义切换、讲义备注、讲义嵌入文字/图片/视频/音频/互动内容/实训内容"),
    ("阅读端", "题库管理与组卷", "多题型题库、自建题库、题目编辑、教材习题导入、组卷、作业发布、考试发布、AI生成试题/试卷"),
    ("阅读端", "资源管理", "学校专业资源库、教材配套资源库、资源分类检索、资源上传下载、在线播放、在线预览、资源共享、权限配置"),
    ("阅读端", "学情督导与精准教学", "学习进度、阅读时长、章节完成情况、任务完成度、教材使用情况、班级作答统计、成绩统计、正确率统计、柱状图分析、学生答题明细"),
    ("阅读端", "分享与打印", "教材二维码分享、章节直达分享、全文打印、按需打印、练习题打印"),
    ("创作者端", "账号与个人设置", "登录、验证码登录、资料维护、头像上传、密码修改、退出账号"),
    ("创作者端", "我的图书管理", "图书列表、图书状态、更新时间、主编信息、新建图书、编辑图书基础资料"),
    ("创作者端", "封面与宣传素材", "封面上传、封面清除、AI生成封面、视频介绍上传、视频介绍清除"),
    ("创作者端", "目录与内容编写", "三级目录创建、编辑、删除、顺序调整、正文撰写、段落样式、字号、字体、对齐、表格、公式、代码块"),
    ("创作者端", "富媒体与互动内容编辑", "图片、图组、音频、视频、附件、网页、习题、案例、实训、背景图、多媒体卡片、互动内容录入与编辑"),
    ("创作者端", "协同创作", "协作者邀请、成员展示、角色说明、成员移除、邀请链接、二维码邀请"),
    ("创作者端", "AI辅助创作", "AI助手、封面生成、内容辅助生成与优化"),
    ("创作者端", "预览、送审与协同修改", "图书预览、保存、提交审核、返修后再提交、审核记录、引用定位、留言回复、讨论线程、版本记录"),
    ("创作者端", "数据统计", "浏览量、任务完成率、视听播放次数、教学模式使用次数统计"),
    ("出版社端", "账号与个人设置", "登录、验证码登录、资料维护、头像上传、密码修改、退出登录"),
    ("出版社端", "稿件状态与教材信息管理", "已出版、待审核、已返修分类展示，主编、版次、ISBN、出版信息管理"),
    ("出版社端", "在线审读", "教材在线阅读、在线审读"),
    ("出版社端", "审核任务处理", "开始审核、查看审核记录、导出审核记录、导出流程记录"),
    ("出版社端", "审核详情与意见管理", "章节目录、正文联动、意见定位、意见卡片展示、意见筛选、标记解决、编辑意见、删除意见、返修意见管理"),
    ("出版社端", "审核流程与结果管理", "审核通过、发起返修、返修详情查看、出版信息查看、已出版教材管理"),
]


FULL_ROWS = [
    ("阅读端", "账号与个人中心", "支持手机号+密码登录、手机号+验证码登录、注册账号、忘记密码、重置密码；支持学校名称、账号名称、姓名、昵称、头像、手机号、密码维护；支持用户反馈提交", "面向学生、教师统一账号体系"),
    ("阅读端", "首页与教材浏览", "支持首页推荐教材展示、教材封面/书名/作者/出版社/ISBN展示、数字出版信息展示、活动展示、友情链接跳转；支持教材搜索、分类浏览、教材详情查看、教材预览", "平台门户与教材发现入口"),
    ("阅读端", "教材获取与书架", "支持在线购买、兑换码激活、免费试读、教材价格展示；支持已购教材管理、教材快捷入口、班级切换、模式快捷跳转", "面向个人教材使用中心"),
    ("阅读端", "智慧阅读", "支持正文阅读、章节目录导航、文内搜索、字号切换、字体切换、阅读背景切换、笔记记录、划线高亮、气泡注释、超链接引用、AI划线解读", "核心数字教材阅读能力"),
    ("阅读端", "富媒体与资源阅读", "支持文字、图片、图组、视频、音频、附件、3D模型等内容嵌入阅读；支持文档、图片、音视频、压缩包、3D模型等多种格式资源下载与部分在线预览", "满足数字教材多媒体展示需求"),
    ("阅读端", "互动答题与AI评测", "支持阅读中嵌入选择题、填空题、判断题、简答题等互动题型；支持自动批阅、答案解析、错题思路、AI综合评价、知识点掌握分析、学习建议生成、重复做题、学习进度追踪", "面向学生学习测评，也支持教师查看结果"),
    ("阅读端", "视听模式", "支持短视频学习、音频学习、学习要点展示、上下一个切换、播放控制、全屏播放", "满足视频化、音频化学习场景"),
    ("阅读端", "任务模式与任务工单", "支持任务式阅读、标准任务、阅读任务、答题任务、实践任务、任务进度跟踪、子任务拆分；支持老师创建任务工单、教材任务导入、附件上传、预览、发布、撤回、任务库复用、任务进度跟踪", "学生用于完成任务，老师角色用于布置和管理任务"),
    ("阅读端", "知识图谱与学情图谱", "支持知识点图谱、知识关联展示、节点点击查看内容、知识点统计；支持知识点掌握、实训掌握、阅读进度、习题进度展示，形成个性化学习路径", "面向学生学习与教师分析"),
    ("阅读端", "实训与交互案例", "支持教材嵌入实训任务、交互案例、实践步骤引导、在线评分评测", "实现理论实践一体化"),
    ("阅读端", "AI伴学", "支持教材专属AI伴学、教材知识问答、相关章节推荐、拓展知识推荐、引导式问答", "面向每本教材的智能辅助学习"),
    ("阅读端", "组群与班级教务", "支持创建群组、加入群组、邀请码/二维码分享、班级教材组织、群组学习统计；支持多班级创建、年级专业配置、班级二维码生成、学生管理、班级信息维护", "学生可参与群组，老师角色可管理班级"),
    ("阅读端", "教学模式", "支持教学模式下展示教材、教案、课件、教学资料预览；支持课件上课、动效播放、翻页教学、全屏上课、教材二分屏展示、分屏大小调节", "老师角色"),
    ("阅读端", "教学辅助工具", "支持分组PK、画笔、多色标注、倒计时、随机抽问、喝彩、课堂积分、小组排行、学生名单导入、AI助教", "老师角色"),
    ("阅读端", "智能云讲义", "支持课时封面展示、幻灯片式讲义切换、讲义备注、讲义中嵌入文字、图片、视频、音频、互动内容、实训内容，并支持课堂讲义辅助授课", "老师角色"),
    ("阅读端", "题库管理与组卷", "支持多题型题库、教师自建题库、按章节/知识点分类、题目编辑删除；支持教材习题导入题库、按知识点/难度/章节组卷、作业发布、考试发布、AI生成试题/试卷", "老师角色"),
    ("阅读端", "资源管理", "支持学校专业资源库、教材配套资源库、资源分类检索、图片/音频/视频/课件/教案/文档上传下载、在线播放、在线预览、资源共享与权限配置", "老师角色"),
    ("阅读端", "学情督导与精准教学", "支持学习进度、阅读时长、章节完成情况、任务完成度、教材使用情况；支持班级人数、作答人数、平均分、最高分、最低分、正确率、逐题正确率、柱状图分析、学生答题明细", "老师角色"),
    ("阅读端", "分享与打印", "支持教材二维码分享、章节直达分享、全文打印、按需打印、练习题打印", "满足传播与线下使用场景"),
    ("创作者端", "账号与个人设置", "支持账号密码登录、手机验证码登录；支持编辑资料、上传头像、设置/修改密码、退出账号", "创作者统一账号体系"),
    ("创作者端", "我的图书管理", "支持查看本人教材列表、教材状态、更新时间、主编信息；支持新建图书、编辑教材名称、封面、作者、简介、视频介绍等基础资料", "创作者工作台首页"),
    ("创作者端", "封面与宣传素材", "支持上传封面、清除封面、AI生成封面、上传视频介绍、清除视频介绍", "教材展示包装能力"),
    ("创作者端", "目录与内容编写", "支持章/节/要点三级目录创建、编辑、删除、顺序调整、展开/收起；支持正文撰写、段落样式、字号、字体、对齐、表格、公式、代码块等基础编辑能力", "核心内容生产能力"),
    ("创作者端", "富媒体与互动内容编辑", "支持图片、图组、音频、视频、附件、网页、习题、案例、实训、背景图、多媒体卡片等内容组件；支持题目互动内容录入、编辑、章节内嵌互动、任务内容配置", "实现教材数字化与互动化"),
    ("创作者端", "协同创作", "支持邀请协作者、成员列表展示、角色说明、成员移除、邀请链接复制、二维码邀请", "满足多人协同编写"),
    ("创作者端", "AI辅助创作", "支持AI助手入口、封面生成、后续教材内容AI辅助生成与优化", "提升创作效率"),
    ("创作者端", "预览、送审与协同修改", "支持图书预览、保存、提交审核、返修后再次提交；支持查看审核记录、引用内容定位、留言回复、展开讨论线程、版本记录查看", "打通创作到审核流程"),
    ("创作者端", "数据统计", "支持浏览量、任务完成率、视听播放次数、教学模式使用次数等数据查看", "为教材优化提供依据"),
    ("创作者端", "个人设置", "支持编辑资料、上传头像、设置/修改密码、退出账号", "创作者账号中心"),
    ("出版社端", "账号与个人设置", "支持账号密码登录、手机验证码登录；支持编辑资料、上传头像、设置/修改密码、退出登录", "出版社统一账号体系"),
    ("出版社端", "稿件状态与教材信息管理", "按已出版、待审核、已返修等状态分类展示教材；展示教材主编、版次、ISBN、出版信息等", "出版社日常管理首页"),
    ("出版社端", "在线审读", "支持直接打开教材内容进行在线审读", "提升审稿效率"),
    ("出版社端", "审核任务处理", "支持开始审核、查看审核记录、导出审核记录、导出流程记录", "满足基础编辑审核流程"),
    ("出版社端", "审核详情与意见管理", "支持章节目录、正文联动、审核意见定位、意见卡片展示；支持审核意见筛选、标记解决、编辑意见、删除意见、返修意见管理", "规范问题闭环处理"),
    ("出版社端", "审核流程与结果管理", "支持审核通过、发起返修、查看返修详情；支持出版信息查看、已出版教材管理", "构成完整审核闭环"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)


def add_title(doc: Document, title: str, subtitle: str):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    p.space_after = Pt(2)

    p2 = doc.add_paragraph()
    r2 = p2.add_run(subtitle)
    r2.font.name = "Arial"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    p2.paragraph_format.space_after = Pt(10)


def add_table(doc: Document, headers, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = text
        cell.width = Cm(widths_cm[i])
        set_cell_shading(cell, "DCE6F1")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.name = "Arial"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                r.font.size = Pt(10)

    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = str(text)
            cells[i].width = Cm(widths_cm[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.name = "Arial"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    r.font.size = Pt(9.5)


def merge_first_column_same_values(table, start_row=1):
    if len(table.rows) <= start_row:
        return
    group_start = start_row
    current_value = table.cell(start_row, 0).text
    for row_idx in range(start_row + 1, len(table.rows)):
        value = table.cell(row_idx, 0).text
        if value != current_value:
            if row_idx - group_start > 1:
                table.cell(group_start, 0).merge(table.cell(row_idx - 1, 0))
            group_start = row_idx
            current_value = value
    if len(table.rows) - group_start > 1:
        table.cell(group_start, 0).merge(table.cell(len(table.rows) - 1, 0))


def create_simple_doc(path: Path):
    doc = Document()
    configure_document(doc)
    add_title(doc, "数字教材平台功能表（简版）", "按阅读端、创作者端、出版社端三类整理，适合汇报或附件概览。")
    add_table(doc, ["端口", "功能模块", "功能点"], SIMPLE_ROWS, [2.2, 4.3, 10.8])
    doc.save(path)


def create_full_doc(path: Path):
    doc = Document()
    configure_document(doc)
    add_title(doc, "数字教材平台功能表（完整版）", "按阅读端、创作者端、出版社端三类整理，包含完整功能描述与说明。")
    add_table(doc, ["端口", "功能模块", "业务功能点", "说明"], FULL_ROWS, [2.0, 3.8, 9.6, 2.8])
    doc.save(path)


def create_full_merged_doc(path: Path):
    doc = Document()
    configure_document(doc)
    add_title(doc, "数字教材平台功能表（完整版）", "按阅读端、创作者端、出版社端三类整理，左侧端口合并展示。")
    rows = [(port, module, details) for port, module, details, _note in FULL_ROWS]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False

    headers = ["端口", "功能模块", "业务功能点"]
    widths_cm = [2.2, 4.1, 11.9]
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = text
        cell.width = Cm(widths_cm[i])
        set_cell_shading(cell, "DCE6F1")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.name = "Arial"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                r.font.size = Pt(10)

    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = str(text)
            cells[i].width = Cm(widths_cm[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.name = "Arial"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    r.font.size = Pt(9.5)

    merge_first_column_same_values(table, start_row=1)
    doc.save(path)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    create_simple_doc(OUTPUT_DIR / "数字教材平台功能表-简版.docx")
    create_full_doc(OUTPUT_DIR / "数字教材平台功能表-完整版.docx")
    create_full_merged_doc(OUTPUT_DIR / "数字教材平台功能表-完整版-合并端口列.docx")


if __name__ == "__main__":
    main()
